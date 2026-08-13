"""
Generates synthetic_test.docx containing synthetic PII, website URLs, and formatting features.
Used to validate PII detection, redaction, formatting preservation, and evaluation metrics.
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_synthetic_document(output_path: str = "examples/synthetic_test.docx"):
    out_file = Path(output_path)
    out_file.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Document Header & Footer
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run("CONFIDENTIAL REPORT — Contact: rashi.patil@gmail.com — Website: www.example-tech.com")
    hrun.font.size = Pt(8)
    hrun.font.color.rgb = RGBColor(128, 128, 128)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fp.add_run("Example Technologies Pvt Ltd — Internal Use Only")
    frun.font.size = Pt(8)

    # Document Title
    title = doc.add_heading(level=0)
    t_run = title.add_run("Synthetic Sample Document for PII Redaction Testing")
    t_run.font.name = "Arial"

    # Paragraph 1: Intro with Name, DOB, Email, Phone, Website
    p1 = doc.add_paragraph()
    p1.add_run("This test document was prepared by ")
    r_name = p1.add_run("Ms. Rashi Patil")
    r_name.bold = True
    p1.add_run(" who was born on ")
    r_dob = p1.add_run("12/04/1998")
    r_dob.italic = True
    p1.add_run(". You can reach her at ")
    r_email = p1.add_run("rashi.patil@gmail.com")
    r_email.bold = True
    r_email.font.color.rgb = RGBColor(0, 51, 102)
    p1.add_run(" or mobile ")
    r_phone = p1.add_run("+91 9876543210")
    r_phone.bold = True
    p1.add_run(" or website ")
    r_url = p1.add_run("https://www.rashi-patil-portfolio.com")
    r_url.underline = True
    p1.add_run(".")

    # Paragraph 2: Promoters (Uppercase names separated by commas!)
    p2 = doc.add_paragraph()
    p2.add_run("OUR PROMOTERS: ")
    p2.add_run("KUSHAL SUBBAYYA HEGDE, PUSHPA KUSHAL HEGDE, RAJESH KUSHAL HEGDE").bold = True
    p2.add_run(".")

    # Paragraph 3: Repeated PII & Secondary Person
    p3 = doc.add_paragraph()
    p3.add_run("Later, ")
    p3.add_run("Ms. Rashi Patil").bold = True
    p3.add_run(" consulted with ")
    p3.add_run("Mr. Rohan Dey").bold = True
    p3.add_run(" regarding the software project at ")
    p3.add_run("Example Technologies Pvt Ltd").bold = True
    p3.add_run(". Mr. Rohan Dey can be reached via ")
    p3.add_run("john.doe@example.com").italic = True
    p3.add_run(" or phone ")
    p3.add_run("+1 555 123 4567")
    p3.add_run(".")

    # Paragraph 4: Address, SSN, Credit Card, IP Address
    p4 = doc.add_paragraph()
    p4.add_run("The official registered address is ")
    p4.add_run("123 Example Road, Sector 17, Chandigarh 160017").underline = True
    p4.add_run(". Primary server IP address is ")
    p4.add_run("192.168.1.10")
    p4.add_run(" with secondary gateway ")
    p4.add_run("10.0.0.1")
    p4.add_run(". Tax identification SSN is ")
    p4.add_run("123-45-6789")
    p4.add_run(" and corporate credit card on file is ")
    p4.add_run("4111 1111 1111 1111")
    p4.add_run(".")

    # Paragraph 5: Table Header / Prospectus Structure Traps
    doc.add_heading("DETAILS OF THE OFFER TO PUBLIC", level=2)
    p5 = doc.add_paragraph()
    p5.add_run("Please process Order number: ")
    p5.add_run("ORD-123456").bold = True
    p5.add_run(" and Ticket number: ")
    p5.add_run("TKT-987654")
    p5.add_run(". Corresponding Invoice number is ")
    p5.add_run("INV-2026-001")
    p5.add_run(" for Product ID ")
    p5.add_run("PROD-12345")
    p5.add_run(". Date of Filing: ")
    p5.add_run("15/08/2026")
    p5.add_run(" with quarterly revenue of ")
    p5.add_run("12,500,000 INR")
    p5.add_run(".")

    # Table with PII and Website column
    doc.add_heading("Employee Registry Table", level=2)
    table = doc.add_table(rows=4, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Employee Name", "Email Address", "Phone Number", "Role / Company", "Website"]
    hdr_cells = table.rows[0].cells
    for i, title_text in enumerate(headers):
        hdr_cells[i].text = title_text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True

    row1 = table.rows[1].cells
    row1[0].text = "Rashi Patil"
    row1[1].text = "rashi.patil@gmail.com"
    row1[2].text = "+91 9876543210"
    row1[3].text = "Example Technologies Pvt Ltd"
    row1[4].text = "www.rashi-patil.com"

    row2 = table.rows[2].cells
    row2[0].text = "Rohan Dey"
    row2[1].text = "rohan.dey@example.org"
    row2[2].text = "+91 9123456789"
    row2[3].text = "Apex Solutions Inc."
    row2[4].text = "www.apexsolutions.com"

    row3 = table.rows[3].cells
    row3[0].text = "Dr. Amit Sharma"
    row3[1].text = "amit.sharma@example.com"
    row3[2].text = "+91-98765-43210"
    row3[3].text = "Global Systems Ltd"
    row3[4].text = "www.globalsystems.co.in"

    doc.save(str(out_file))
    print(f"Synthetic test document generated successfully at: {out_file.resolve()}")


if __name__ == "__main__":
    create_synthetic_document()
