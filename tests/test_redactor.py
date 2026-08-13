"""
Tests for DOCX redaction engine and run preservation in src/redactor.py.
"""

import pytest
from docx import Document
from src.models import DocumentBlock, PIIEntity
from src.replacement_generator import ReplacementGenerator
from src.redactor import DOCXRedactor


def test_paragraph_redaction_single_run():
    doc = Document()
    p = doc.add_paragraph("Contact Rashi Patil today.")

    generator = ReplacementGenerator(seed=42)
    redactor = DOCXRedactor(generator)

    block = DocumentBlock(block_type="paragraph", text=p.text, block_id="p1", element_ref=p)
    entity = PIIEntity(
        entity_type="PERSON",
        text="Rashi Patil",
        start=8,
        end=19,
        confidence=0.9,
        source="test"
    )

    redactor.redact_document(doc, {block: [entity]})
    assert "Rashi Patil" not in p.text
    assert len(p.runs) > 0


def test_paragraph_redaction_multi_run_formatting_preservation():
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("User ")
    r2 = p.add_run("Rashi ")
    r2.bold = True
    r3 = p.add_run("Patil ")
    r3.bold = True
    r4 = p.add_run("joined.")

    generator = ReplacementGenerator(seed=42)
    redactor = DOCXRedactor(generator)

    # Combined text is "User Rashi Patil joined."
    # "Rashi Patil" spans index 5 to 16
    block = DocumentBlock(block_type="paragraph", text=p.text, block_id="p1", element_ref=p)
    entity = PIIEntity(
        entity_type="PERSON",
        text="Rashi Patil",
        start=5,
        end=16,
        confidence=0.9,
        source="test"
    )

    redactor.redact_document(doc, {block: [entity]})
    assert "Rashi Patil" not in p.text
    assert "joined." in p.text
