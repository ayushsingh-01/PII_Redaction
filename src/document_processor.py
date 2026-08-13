"""
Document Processor for extracting logical text blocks from DOCX files.
Processes Paragraphs, Table Cells, Headers, and Footers.
"""

import logging
from typing import List, Tuple
from docx import Document
from docx.text.paragraph import Paragraph

from src.models import DocumentBlock

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Extracts text blocks from python-docx Document objects.
    """
    def __init__(self, docx_path: str):
        self.docx_path = docx_path
        try:
            self.doc = Document(docx_path)
        except Exception as e:
            logger.error(f"Failed to open DOCX file {docx_path}: {e}")
            raise e

    def extract_blocks(self) -> List[DocumentBlock]:
        """
        Extracts all editable text blocks (paragraphs, tables, headers, footers).
        """
        blocks: List[DocumentBlock] = []
        block_counter = 0

        # 1. Main Document Paragraphs
        for p in self.doc.paragraphs:
            if p.text:
                block_counter += 1
                blocks.append(DocumentBlock(
                    block_type="paragraph",
                    text=p.text,
                    block_id=f"p_{block_counter}",
                    element_ref=p
                ))

        # 2. Main Document Tables
        for t_idx, table in enumerate(self.doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    for p in cell.paragraphs:
                        if p.text:
                            block_counter += 1
                            blocks.append(DocumentBlock(
                                block_type="table_cell",
                                text=p.text,
                                block_id=f"t{t_idx}_r{r_idx}_c{c_idx}_p{block_counter}",
                                element_ref=p
                            ))

        # 3. Section Headers & Footers
        for s_idx, section in enumerate(self.doc.sections):
            # Header Paragraphs & Tables
            if section.header and not section.header.is_linked_to_previous:
                for p in section.header.paragraphs:
                    if p.text:
                        block_counter += 1
                        blocks.append(DocumentBlock(
                            block_type="header",
                            text=p.text,
                            block_id=f"sec{s_idx}_hdr_p{block_counter}",
                            element_ref=p,
                            section_index=s_idx
                        ))
                for t_idx, table in enumerate(section.header.tables):
                    for r_idx, row in enumerate(table.rows):
                        for c_idx, cell in enumerate(row.cells):
                            for p in cell.paragraphs:
                                if p.text:
                                    block_counter += 1
                                    blocks.append(DocumentBlock(
                                        block_type="header",
                                        text=p.text,
                                        block_id=f"sec{s_idx}_hdr_t{t_idx}_p{block_counter}",
                                        element_ref=p,
                                        section_index=s_idx
                                    ))

            # Footer Paragraphs & Tables
            if section.footer and not section.footer.is_linked_to_previous:
                for p in section.footer.paragraphs:
                    if p.text:
                        block_counter += 1
                        blocks.append(DocumentBlock(
                            block_type="footer",
                            text=p.text,
                            block_id=f"sec{s_idx}_ftr_p{block_counter}",
                            element_ref=p,
                            section_index=s_idx
                        ))
                for t_idx, table in enumerate(section.footer.tables):
                    for r_idx, row in enumerate(table.rows):
                        for c_idx, cell in enumerate(row.cells):
                            for p in cell.paragraphs:
                                if p.text:
                                    block_counter += 1
                                    blocks.append(DocumentBlock(
                                        block_type="footer",
                                        text=p.text,
                                        block_id=f"sec{s_idx}_ftr_t{t_idx}_p{block_counter}",
                                        element_ref=p,
                                        section_index=s_idx
                                    ))

        return blocks
